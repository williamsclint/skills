import os
import re
import io
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Wynn brand colours ────────────────────────────────────────────────────────
GOLD      = RGBColor(0xB8, 0x96, 0x55)   # Wynn gold
DARK_NAVY = RGBColor(0x1A, 0x1A, 0x2E)   # deep navy
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)
MID_GREY  = RGBColor(0x6B, 0x6B, 0x6B)
BODY_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
RED_FLAG  = RGBColor(0xC0, 0x39, 0x2B)

LOGO_URL = "https://images.ctfassets.net/nmhu6zvnbnyf/6OGKMyj2vnB0ROqomayO0E/6e468d63e0f43d47b3c94c64e92ad1ee/image.png"

FONT_HEADING = "Garamond"
FONT_BODY    = "Garamond"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_background(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc, color: RGBColor = GOLD, width_pt: int = 2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(width_pt * 8))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*color))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def run_font(run, name, size_pt, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def para_font(para, name, size_pt, bold=False, italic=False,
              color=None, align=None, space_before=0, space_after=4):
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    if align:
        para.alignment = align
    for run in para.runs:
        run_font(run, name, size_pt, bold, italic, color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    if level == 1:
        run_font(run, FONT_HEADING, 15, bold=True, color=DARK_NAVY)
        add_horizontal_rule(doc, GOLD, 2)
    elif level == 2:
        run_font(run, FONT_HEADING, 12, bold=True, color=DARK_NAVY)
        add_horizontal_rule(doc, MID_GREY, 1)
    else:
        run_font(run, FONT_HEADING, 11, bold=True, italic=True, color=DARK_NAVY)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    # Handle inline bold (**text**)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run_font(run, FONT_BODY, 10, bold=True, color=BODY_TEXT)
        else:
            run = p.add_run(part)
            run_font(run, FONT_BODY, 10, color=BODY_TEXT)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run_font(run, FONT_BODY, 10, bold=True, color=BODY_TEXT)
        else:
            run = p.add_run(part)
            run_font(run, FONT_BODY, 10, color=BODY_TEXT)
    return p

def add_markdown_table(doc, header_row, data_rows):
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True

    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header_row):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1A1A2E")
        for run in hdr_cells[i].paragraphs[0].runs:
            run_font(run, FONT_BODY, 9, bold=True, color=WHITE)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(data_rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = "F5F5F5" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row):
            row_cells[c_idx].text = cell_text
            set_cell_background(row_cells[c_idx], bg)
            for run in row_cells[c_idx].paragraphs[0].runs:
                run_font(run, FONT_BODY, 9, color=BODY_TEXT)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

# ── Letterhead ────────────────────────────────────────────────────────────────

def add_letterhead(doc, logo_bytes, process_title, confidential_label="CONFIDENTIAL"):
    # Top accent bar via a 1-row, 1-col table
    bar = doc.add_table(rows=1, cols=1)
    bar.autofit = False
    bar.columns[0].width = Inches(7.5)
    cell = bar.rows[0].cells[0]
    cell.height = Cm(0.35)
    set_cell_background(cell, "B89655")
    cell.text = ""
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Logo + title side by side
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.autofit = False
    header_tbl.columns[0].width = Inches(2.0)
    header_tbl.columns[1].width = Inches(5.5)

    # Logo cell
    logo_cell = header_tbl.rows[0].cells[0]
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = logo_para.add_run()
    logo_run.add_picture(io.BytesIO(logo_bytes), width=Inches(1.7))

    # Title cell
    title_cell = header_tbl.rows[0].cells[1]
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_brand = title_cell.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_brand = p_brand.add_run("WYNN RESORTS")
    run_font(r_brand, FONT_HEADING, 9, bold=True, color=GOLD)
    p_brand.paragraph_format.space_after = Pt(1)

    p_dept = title_cell.add_paragraph()
    p_dept.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_dept = p_dept.add_run("Office of the COO  |  Operations & Compliance")
    run_font(r_dept, FONT_BODY, 8, italic=True, color=MID_GREY)
    p_dept.paragraph_format.space_after = Pt(2)

    p_conf = title_cell.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_conf = p_conf.add_run(confidential_label)
    run_font(r_conf, FONT_BODY, 7.5, bold=True, color=RED_FLAG)
    p_conf.paragraph_format.space_after = Pt(0)

    # Thin gold rule
    add_horizontal_rule(doc, GOLD, 2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Document title block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_label = p_title.add_run("PROCESS AUDIT INTAKE DOCUMENT\n")
    run_font(r_label, FONT_HEADING, 9, bold=True, color=MID_GREY)
    r_title = p_title.add_run(process_title)
    run_font(r_title, FONT_HEADING, 18, bold=True, color=DARK_NAVY)
    p_title.paragraph_format.space_after = Pt(4)

    add_horizontal_rule(doc, GOLD, 1)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Metadata box ──────────────────────────────────────────────────────────────

def add_meta_box(doc, fields: dict):
    """Renders a 2-col key-value metadata table with navy header strip."""
    items = list(fields.items())
    table = doc.add_table(rows=len(items), cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(5.5)
    table.style = "Table Grid"

    for i, (key, val) in enumerate(items):
        key_cell = table.rows[i].cells[0]
        val_cell = table.rows[i].cells[1]
        bg = "1A1A2E" if i == 0 else ("F5F5F5" if i % 2 == 0 else "FFFFFF")
        set_cell_background(key_cell, bg)
        set_cell_background(val_cell, bg)
        key_cell.text = key
        val_cell.text = val
        kc = WHITE if i == 0 else DARK_NAVY
        vc = WHITE if i == 0 else BODY_TEXT
        for run in key_cell.paragraphs[0].runs:
            run_font(run, FONT_BODY, 9, bold=True, color=kc)
        for run in val_cell.paragraphs[0].runs:
            run_font(run, FONT_BODY, 9, color=vc)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── Footer ────────────────────────────────────────────────────────────────────

def add_footer(doc, filename):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Wynn Resorts, Limited  |  Internal Use Only  |  {filename}  "
        f"|  Do not distribute without authorization"
    )
    run_font(run, FONT_BODY, 7.5, italic=True, color=MID_GREY)

# ── Margin setup ─────────────────────────────────────────────────────────────

def set_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)

# ── Markdown parser → doc ────────────────────────────────────────────────────

def is_table_line(line):
    return line.strip().startswith("|") and "|" in line[1:]

def parse_table_block(lines):
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-| :]+\|\s*$", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows

def render_markdown_to_doc(doc, md_text):
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Table block
        if is_table_line(line):
            tbl_lines = []
            while i < len(lines) and is_table_line(lines[i]):
                tbl_lines.append(lines[i])
                i += 1
            rows = parse_table_block(tbl_lines)
            if rows:
                add_markdown_table(doc, rows[0], rows[1:])
            continue

        stripped = line.strip()

        # Section 1 heading (## Section N —)
        if stripped.startswith("## Section"):
            add_heading(doc, stripped.lstrip("# "), level=1)

        # H1
        elif stripped.startswith("# "):
            pass  # title already in letterhead

        # H2
        elif stripped.startswith("## "):
            add_heading(doc, stripped.lstrip("# "), level=2)

        # H3 / H4
        elif stripped.startswith("### ") or stripped.startswith("#### "):
            add_heading(doc, stripped.lstrip("# "), level=3)

        # Bullet
        elif stripped.startswith("- "):
            add_bullet(doc, stripped[2:])

        # Numbered list
        elif re.match(r"^\d+\. ", stripped):
            add_bullet(doc, re.sub(r"^\d+\. ", "", stripped))

        # Horizontal rule
        elif stripped.startswith("---"):
            add_horizontal_rule(doc, GOLD, 1)

        # Bold-label paragraph (e.g. **Process to audit:**)
        elif stripped.startswith("**") and stripped.endswith("**"):
            add_body(doc, stripped)

        # Non-empty body text
        elif stripped:
            add_body(doc, stripped)

        i += 1

# ── Per-document metadata extraction ─────────────────────────────────────────

def extract_meta(md_text: str) -> dict:
    """Pull key-value pairs from Section 1 of the markdown."""
    meta = {}
    in_section1 = False
    for line in md_text.splitlines():
        if "Section 1" in line:
            in_section1 = True
            continue
        if in_section1 and line.startswith("## Section"):
            break
        if in_section1:
            m = re.match(r"-\s+\*\*(.+?):\*\*\s+(.*)", line)
            if m:
                meta[m.group(1)] = m.group(2)
    return meta

# ── Document definitions ─────────────────────────────────────────────────────

DOCS = [
    {
        "md":    "01-vip-credit-issuance.md",
        "out":   "01-VIP-Casino-Credit-Issuance.docx",
        "title": "VIP Patron Casino Credit Issuance",
    },
    {
        "md":    "02-gaming-table-fill-credit.md",
        "out":   "02-Gaming-Table-Fill-Credit.docx",
        "title": "Gaming Table Chip Fill & Credit Process",
    },
    {
        "md":    "03-gaming-license-onboarding.md",
        "out":   "03-Gaming-License-Onboarding.docx",
        "title": "New Gaming Employee Licensing & Onboarding",
    },
    {
        "md":    "04-sar-ctr-compliance-filing.md",
        "out":   "04-BSA-AML-SAR-CTR-Filing.docx",
        "title": "BSA/AML SAR & CTR Filing Process",
    },
    {
        "md":    "05-fb-vendor-onboarding.md",
        "out":   "05-Fine-Dining-FB-Vendor-Onboarding.docx",
        "title": "Fine Dining F&B Vendor Onboarding",
    },
]

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    base = os.path.dirname(os.path.abspath(__file__))

    print("Downloading Wynn logo…")
    resp = requests.get(LOGO_URL, timeout=15)
    resp.raise_for_status()
    logo_bytes = resp.content
    print(f"  Logo downloaded ({len(logo_bytes):,} bytes)")

    for doc_def in DOCS:
        md_path  = os.path.join(base, doc_def["md"])
        out_path = os.path.join(base, doc_def["out"])

        print(f"\nBuilding: {doc_def['out']}")
        with open(md_path, encoding="utf-8") as f:
            md_text = f.read()

        doc = Document()
        set_margins(doc, top=0.85, bottom=0.85, left=1.0, right=1.0)

        add_letterhead(doc, logo_bytes, doc_def["title"])

        # Metadata summary box from Section 1
        meta = extract_meta(md_text)
        if meta:
            box_fields = {"DOCUMENT METADATA": ""}
            box_fields.update({k: v for k, v in list(meta.items())[:6]})
            add_meta_box(doc, box_fields)

        # Render the rest of the markdown (skip H1 title line)
        body_md = "\n".join(
            l for l in md_text.splitlines()
            if not l.startswith("# Process Audit Intake:")
        )
        render_markdown_to_doc(doc, body_md)

        add_footer(doc, doc_def["out"])
        doc.save(out_path)
        print(f"  Saved -> {out_path}")

    print("\nDone. All 5 documents generated.")

if __name__ == "__main__":
    main()
