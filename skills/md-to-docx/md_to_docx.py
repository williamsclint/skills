"""
md_to_docx.py — Convert a markdown file to a formatted Word document.
Usage: python md_to_docx.py <input.md> <output.docx>
"""
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("python-docx not installed. Run: python -m pip install python-docx")
    sys.exit(1)

DARK_BLUE = RGBColor(0, 51, 102)
BODY_COLOR = RGBColor(26, 26, 26)

def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_COLOR
    style.paragraph_format.space_after = Pt(6)
    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = DARK_BLUE
        hs.font.name = 'Calibri'
        sizes = {1: 16, 2: 14, 3: 12}
        hs.font.size = Pt(sizes[level])
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            if c_idx < len(table.columns):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = val.strip()
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
    doc.add_paragraph()

def add_rich_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    bold_pattern = re.compile(r'\*\*(.+?)\*\*')
    italic_pattern = re.compile(r'\*(.+?)\*')
    parts = bold_pattern.split(text)
    for i, part in enumerate(parts):
        if i % 2 == 1:
            run = p.add_run(part)
            run.bold = True
        else:
            italic_parts = italic_pattern.split(part)
            for j, ipart in enumerate(italic_parts):
                if j % 2 == 1:
                    run = p.add_run(ipart)
                    run.italic = True
                elif ipart:
                    p.add_run(ipart)
    return p

def convert(input_path, output_path):
    md = Path(input_path).read_text(encoding='utf-8')
    lines = md.split('\n')
    doc = Document()
    setup_styles(doc)

    i = 0
    in_code_block = False
    in_table = False
    table_headers = []
    table_rows = []
    meta_lines = []
    title_done = False

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue
        if in_code_block:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            i += 1
            continue

        # Flush table if we were in one and this line isn't a table row
        if in_table and not line.strip().startswith('|'):
            add_styled_table(doc, table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []

        # Table rows
        if line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if not in_table:
                table_headers = cells
                in_table = True
            elif all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
                pass  # separator row
            else:
                table_rows.append(cells)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            if not title_done:
                i += 1
                continue
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if level == 1 and not title_done:
                h = doc.add_heading(text, level=0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_done = True
            else:
                doc.add_heading(text, level=level)
            i += 1
            continue

        # Metadata lines (bold key: value)
        meta_match = re.match(r'^\*\*(.+?):\*\*\s*(.+)$', line)
        if meta_match and not title_done or (title_done and len(doc.paragraphs) < 5):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{meta_match.group(1)}: ")
            run.bold = True
            p.add_run(meta_match.group(2))
            i += 1
            continue

        # Blockquotes
        if line.strip().startswith('>'):
            text = line.strip().lstrip('>').strip()
            p = add_rich_paragraph(doc, text)
            p.paragraph_format.left_indent = Cm(1.5)
            p.runs[0].italic = True if p.runs else None
            i += 1
            continue

        # Bullet lists
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if bullet_match:
            text = bullet_match.group(2)
            add_rich_paragraph(doc, text, style='List Bullet')
            i += 1
            continue

        # Sub-bullets (indented)
        sub_bullet_match = re.match(r'^\s{2,}[-*]\s+(.+)$', line)
        if sub_bullet_match:
            text = sub_bullet_match.group(1)
            add_rich_paragraph(doc, text, style='List Bullet 2')
            i += 1
            continue

        # Regular paragraph
        if line.strip():
            add_rich_paragraph(doc, line.strip())

        i += 1

    # Flush any remaining table
    if in_table:
        add_styled_table(doc, table_headers, table_rows)

    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
