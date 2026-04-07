"""
wynn_doc_utils.py
Shared utility library for generating Wynn Resorts branded .docx documents.
Implements the standards defined in wynn-doc-format/SKILL.md.
"""

import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand palette ─────────────────────────────────────────────────────────────
GOLD        = RGBColor(0xB8, 0x96, 0x55)
DARK_NAVY   = RGBColor(0x1A, 0x1A, 0x2E)
BODY_TEXT   = RGBColor(0x1A, 0x1A, 0x1A)
MID_GREY    = RGBColor(0x6B, 0x6B, 0x6B)
LIGHT_GREY  = RGBColor(0xF5, 0xF5, 0xF5)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
ALERT_RED   = RGBColor(0xC0, 0x39, 0x2B)
ALERT_ORANGE= RGBColor(0xE6, 0x7E, 0x22)
ALERT_AMBER = RGBColor(0xF3, 0x9C, 0x12)
ALERT_GREEN = RGBColor(0x27, 0xAE, 0x60)

RISK_COLORS = {
    "CRITICAL": ("C0392B", "CRITICAL"),
    "HIGH":     ("E67E22", "HIGH"),
    "MEDIUM":   ("F39C12", "MEDIUM"),
    "LOW":      ("27AE60", "LOW"),
}

FONT = "Garamond"


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _hex(rgb: RGBColor) -> str:
    return "{:02X}{:02X}{:02X}".format(*rgb)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)


def _run(run, size, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _para_spacing(para, before=0, after=4):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


def _inline_bold(para, text: str, size: float, base_color=None):
    """Render **bold** markers in text as bold runs."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            _run(r, size, bold=True, color=base_color)
        else:
            r = para.add_run(part)
            _run(r, size, color=base_color)


# ── Public primitives ─────────────────────────────────────────────────────────

def new_doc() -> Document:
    return Document()


def set_margins(doc, top=0.85, bottom=0.85, left=1.0, right=1.0):
    for s in doc.sections:
        s.top_margin    = Inches(top)
        s.bottom_margin = Inches(bottom)
        s.left_margin   = Inches(left)
        s.right_margin  = Inches(right)


def add_horizontal_rule(doc, color: RGBColor = GOLD, weight: int = 2):
    p = doc.add_paragraph()
    _para_spacing(p, 2, 2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(weight * 8))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), _hex(color))
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    _para_spacing(p, 14 if level == 1 else 8, 4)
    r = p.add_run(text)
    if level == 1:
        _run(r, 15, bold=True, color=DARK_NAVY)
        add_horizontal_rule(doc, GOLD, 2)
    elif level == 2:
        _run(r, 12, bold=True, color=DARK_NAVY)
        add_horizontal_rule(doc, MID_GREY, 1)
    else:
        _run(r, 11, bold=True, italic=True, color=DARK_NAVY)
    return p


def add_body(doc, text: str, indent=False):
    p = doc.add_paragraph()
    _para_spacing(p, 1, 4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    _inline_bold(p, text, 10, BODY_TEXT)
    return p


def add_bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    _para_spacing(p, 1, 2)
    _inline_bold(p, text, 10, BODY_TEXT)
    return p


def add_spacer(doc, pts: int = 4):
    p = doc.add_paragraph()
    _para_spacing(p, 0, pts)
    return p


def add_markdown_table(doc, header: list, rows: list):
    if not rows:
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Table Grid"
    tbl.autofit = True
    hcells = tbl.rows[0].cells
    for i, h in enumerate(header):
        hcells[i].text = h
        set_cell_bg(hcells[i], "1A1A2E")
        for run in hcells[i].paragraphs[0].runs:
            _run(run, 9, bold=True, color=WHITE)
        hcells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        bg = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        dcells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row):
            dcells[ci].text = str(val)
            set_cell_bg(dcells[ci], bg)
            for run in dcells[ci].paragraphs[0].runs:
                _run(run, 9, color=BODY_TEXT)
    add_spacer(doc, 6)
    return tbl


def add_meta_box(doc, fields: dict):
    items = list(fields.items())
    tbl = doc.add_table(rows=len(items), cols=2)
    tbl.style = "Table Grid"
    tbl.autofit = False
    tbl.columns[0].width = Inches(2.0)
    tbl.columns[1].width = Inches(5.5)
    for i, (k, v) in enumerate(items):
        kc, vc = tbl.rows[i].cells[0], tbl.rows[i].cells[1]
        bg = "1A1A2E" if i == 0 else ("F5F5F5" if i % 2 == 0 else "FFFFFF")
        set_cell_bg(kc, bg); set_cell_bg(vc, bg)
        kc.text = k; vc.text = str(v)
        tc = WHITE if i == 0 else DARK_NAVY
        vc_color = WHITE if i == 0 else BODY_TEXT
        for run in kc.paragraphs[0].runs:
            _run(run, 9, bold=True, color=tc)
        for run in vc.paragraphs[0].runs:
            _run(run, 9, color=vc_color)
    add_spacer(doc, 8)


# ── Letterhead ────────────────────────────────────────────────────────────────

def add_letterhead(doc, logo_bytes: bytes, title: str,
                   department: str = "Office of the COO  |  Operations & Compliance",
                   classification: str = "STRICTLY CONFIDENTIAL"):
    # Gold accent bar
    bar = doc.add_table(rows=1, cols=1)
    bar.autofit = False
    bar.columns[0].width = Inches(7.5)
    cell = bar.rows[0].cells[0]
    set_cell_bg(cell, "B89655")
    cell.text = ""
    add_spacer(doc, 3)

    # Header: logo | brand info
    htbl = doc.add_table(rows=1, cols=2)
    htbl.autofit = False
    htbl.columns[0].width = Inches(2.0)
    htbl.columns[1].width = Inches(5.5)

    lc = htbl.rows[0].cells[0]
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.add_run().add_picture(io.BytesIO(logo_bytes), width=Inches(1.7))

    rc = htbl.rows[0].cells[1]
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p1 = rc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _para_spacing(p1, 0, 1)
    _run(p1.add_run("WYNN RESORTS"), 9, bold=True, color=GOLD)

    p2 = rc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _para_spacing(p2, 0, 2)
    _run(p2.add_run(department), 8, italic=True, color=MID_GREY)

    p3 = rc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _para_spacing(p3, 0, 0)
    _run(p3.add_run(classification), 7.5, bold=True, color=ALERT_RED)

    add_horizontal_rule(doc, GOLD, 2)
    add_spacer(doc, 3)

    # Title block
    pt = doc.add_paragraph()
    _para_spacing(pt, 0, 4)
    _run(pt.add_run("INTERNAL AUDIT REPORT\n"), 9, bold=True, color=MID_GREY)
    _run(pt.add_run(title), 18, bold=True, color=DARK_NAVY)

    add_horizontal_rule(doc, GOLD, 1)
    add_spacer(doc, 6)


def add_footer(doc, filename: str):
    sec = doc.sections[0]
    footer = sec.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"Wynn Resorts, Limited  |  Strictly Confidential  |  {filename}  "
        f"|  Do not distribute without authorization"
    )
    _run(r, 7.5, italic=True, color=MID_GREY)


# ── Audit-specific components ─────────────────────────────────────────────────

def add_exec_summary_block(doc, overall_rating: str, counts: dict, bullets: list):
    """
    overall_rating: 'Satisfactory' | 'Needs Improvement' | 'Unsatisfactory'
    counts: {'Critical': n, 'High': n, 'Medium': n, 'Low': n}
    bullets: list of strings
    """
    add_heading(doc, "Executive Summary", 1)

    # Overall rating badge
    rating_colors = {
        "Satisfactory":      ("27AE60", WHITE),
        "Needs Improvement": ("F39C12", WHITE),
        "Unsatisfactory":    ("C0392B", WHITE),
    }
    bg, fg = rating_colors.get(overall_rating, ("6B6B6B", WHITE))

    rtbl = doc.add_table(rows=1, cols=2)
    rtbl.autofit = False
    rtbl.columns[0].width = Inches(1.5)
    rtbl.columns[1].width = Inches(6.0)
    rc_label = rtbl.rows[0].cells[0]
    rc_value = rtbl.rows[0].cells[1]
    set_cell_bg(rc_label, "1A1A2E")
    set_cell_bg(rc_value, bg)
    rc_label.text = "OVERALL RATING"
    rc_value.text = overall_rating.upper()
    for run in rc_label.paragraphs[0].runs:
        _run(run, 8, bold=True, color=WHITE)
    for run in rc_value.paragraphs[0].runs:
        _run(run, 10, bold=True, color=WHITE)
    add_spacer(doc, 4)

    # Finding count summary
    ctbl = doc.add_table(rows=2, cols=4)
    ctbl.style = "Table Grid"
    ctbl.autofit = False
    for i in range(4): ctbl.columns[i].width = Inches(1.875)
    labels = ["CRITICAL", "HIGH", "MEDIUM", "LOW"]
    for i, label in enumerate(labels):
        hc = ctbl.rows[0].cells[i]
        vc = ctbl.rows[1].cells[i]
        set_cell_bg(hc, RISK_COLORS[label][0])
        set_cell_bg(vc, "FFFFFF")
        hc.text = label
        vc.text = str(counts.get(label.capitalize(), counts.get(label, 0)))
        for run in hc.paragraphs[0].runs:
            _run(run, 9, bold=True, color=WHITE)
            hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in vc.paragraphs[0].runs:
            _run(run, 14, bold=True, color=DARK_NAVY)
            vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_spacer(doc, 6)

    # Key findings bullets
    p = doc.add_paragraph()
    _para_spacing(p, 0, 4)
    _run(p.add_run("Key Findings"), 11, bold=True, color=DARK_NAVY)
    for bullet in bullets:
        add_bullet(doc, bullet)
    add_spacer(doc, 4)


def add_scope_methodology(doc, scope: str, methodology: str):
    add_heading(doc, "Audit Scope & Methodology", 1)
    add_body(doc, "**Scope:** " + scope)
    add_spacer(doc, 2)
    add_body(doc, "**Methodology:** " + methodology)
    add_spacer(doc, 4)


def add_findings_summary_table(doc, findings: list):
    add_heading(doc, "Summary of Findings", 1)
    header = ["ID", "Finding Title", "Risk Rating", "Owner", "Target"]
    rows = []
    for f in findings:
        rows.append([
            f["id"], f["title"], f["risk"],
            f.get("owner", "—"), f.get("target", "—")
        ])
    add_markdown_table(doc, header, rows)


def add_finding_block(doc, finding: dict):
    fid   = finding["id"]
    title = finding["title"]
    risk  = finding["risk"].upper()
    risk_hex, _ = RISK_COLORS.get(risk, ("6B6B6B", "—"))

    # Finding header bar
    htbl = doc.add_table(rows=1, cols=3)
    htbl.autofit = False
    htbl.columns[0].width = Inches(0.75)
    htbl.columns[1].width = Inches(5.5)
    htbl.columns[2].width = Inches(1.25)
    set_cell_bg(htbl.rows[0].cells[0], "1A1A2E")
    set_cell_bg(htbl.rows[0].cells[1], "1A1A2E")
    set_cell_bg(htbl.rows[0].cells[2], risk_hex)

    htbl.rows[0].cells[0].text = fid
    htbl.rows[0].cells[1].text = title
    htbl.rows[0].cells[2].text = risk

    for cell in htbl.rows[0].cells[:2]:
        for run in cell.paragraphs[0].runs:
            _run(run, 10, bold=True, color=WHITE)
    for run in htbl.rows[0].cells[2].paragraphs[0].runs:
        _run(run, 9, bold=True, color=WHITE)
    htbl.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Detail grid: condition / criteria / cause / effect
    fields = [
        ("CONDITION",      finding.get("condition", "")),
        ("CRITERIA",       finding.get("criteria",  "")),
        ("CAUSE",          finding.get("cause",     "")),
        ("EFFECT",         finding.get("effect",    "")),
    ]
    for label, content in fields:
        dtbl = doc.add_table(rows=1, cols=2)
        dtbl.style = "Table Grid"
        dtbl.autofit = False
        dtbl.columns[0].width = Inches(1.25)
        dtbl.columns[1].width = Inches(6.25)
        lc, vc = dtbl.rows[0].cells[0], dtbl.rows[0].cells[1]
        set_cell_bg(lc, "1A1A2E")
        set_cell_bg(vc, "FAFAFA")
        lc.text = label
        lc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for run in lc.paragraphs[0].runs:
            _run(run, 8, bold=True, color=WHITE)
        vc.text = content
        for run in vc.paragraphs[0].runs:
            _run(run, 9, color=BODY_TEXT)

    # Recommendations box
    if finding.get("recommendations"):
        rtbl = doc.add_table(rows=1, cols=1)
        rtbl.style = "Table Grid"
        rtbl.autofit = False
        rtbl.columns[0].width = Inches(7.5)
        rc = rtbl.rows[0].cells[0]
        set_cell_bg(rc, "FBF6E9")
        rp = rc.paragraphs[0]
        _run(rp.add_run("RECOMMENDATIONS"), 8, bold=True, color=DARK_NAVY)
        _para_spacing(rp, 2, 2)
        for rec in finding["recommendations"]:
            rp2 = rc.add_paragraph()
            _para_spacing(rp2, 1, 2)
            rp2.paragraph_format.left_indent = Inches(0.15)
            _run(rp2.add_run(f"  {chr(0x2022)}  {rec}"), 9, color=BODY_TEXT)

    # Owner / Target / Management response row
    otbl = doc.add_table(rows=2, cols=3)
    otbl.style = "Table Grid"
    otbl.autofit = False
    otbl.columns[0].width = Inches(2.5)
    otbl.columns[1].width = Inches(2.5)
    otbl.columns[2].width = Inches(2.5)

    labels2 = ["RESPONSIBLE OWNER", "TARGET DATE", "MANAGEMENT RESPONSE"]
    values2 = [
        finding.get("owner", "—"),
        finding.get("target", "—"),
        finding.get("mgmt_response", "Response pending — due within 30 days of report issuance."),
    ]
    for i, (lbl, val) in enumerate(zip(labels2, values2)):
        lc2 = otbl.rows[0].cells[i]
        vc2 = otbl.rows[1].cells[i]
        set_cell_bg(lc2, "1A1A2E")
        set_cell_bg(vc2, "F5F5F5")
        lc2.text = lbl
        vc2.text = val
        for run in lc2.paragraphs[0].runs:
            _run(run, 8, bold=True, color=WHITE)
        for run in vc2.paragraphs[0].runs:
            _run(run, 9, color=BODY_TEXT)

    add_spacer(doc, 10)


def add_observation_block(doc, obs: dict):
    oid   = obs["id"]
    title = obs["title"]
    risk  = obs.get("risk", "LOW").upper()
    risk_hex, _ = RISK_COLORS.get(risk, ("27AE60", "—"))

    otbl = doc.add_table(rows=1, cols=3)
    otbl.autofit = False
    otbl.columns[0].width = Inches(0.75)
    otbl.columns[1].width = Inches(5.5)
    otbl.columns[2].width = Inches(1.25)
    set_cell_bg(otbl.rows[0].cells[0], "6B6B6B")
    set_cell_bg(otbl.rows[0].cells[1], "6B6B6B")
    set_cell_bg(otbl.rows[0].cells[2], risk_hex)
    otbl.rows[0].cells[0].text = oid
    otbl.rows[0].cells[1].text = title
    otbl.rows[0].cells[2].text = risk
    for cell in otbl.rows[0].cells[:2]:
        for run in cell.paragraphs[0].runs:
            _run(run, 10, bold=True, color=WHITE)
    for run in otbl.rows[0].cells[2].paragraphs[0].runs:
        _run(run, 9, bold=True, color=WHITE)
    otbl.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    dtbl = doc.add_table(rows=1, cols=1)
    dtbl.style = "Table Grid"
    dtbl.columns[0].width = Inches(7.5)
    set_cell_bg(dtbl.rows[0].cells[0], "FAFAFA")
    dtbl.rows[0].cells[0].text = obs.get("detail", "")
    for run in dtbl.rows[0].cells[0].paragraphs[0].runs:
        _run(run, 9, color=BODY_TEXT)

    if obs.get("recommendation"):
        rtbl = doc.add_table(rows=1, cols=1)
        rtbl.style = "Table Grid"
        rtbl.columns[0].width = Inches(7.5)
        rc = rtbl.rows[0].cells[0]
        set_cell_bg(rc, "FBF6E9")
        rp = rc.paragraphs[0]
        _run(rp.add_run("RECOMMENDATION:  "), 8, bold=True, color=DARK_NAVY)
        _run(rp.add_run(obs["recommendation"]), 9, color=BODY_TEXT)

    add_spacer(doc, 8)


def add_prior_findings_table(doc, prior: list):
    """prior: list of dicts with keys: ref, title, original_rating, status, notes"""
    add_heading(doc, "Prior Audit Findings — Status Update", 1)
    if not prior:
        add_body(doc, "No prior audit findings applicable to this process.")
        return
    header = ["Ref", "Finding", "Original Rating", "Current Status", "Notes"]
    rows = [[p["ref"], p["title"], p["original_rating"], p["status"], p.get("notes", "")] for p in prior]
    add_markdown_table(doc, header, rows)


def add_sign_off_block(doc, audit_team: list, distribution: list, report_date: str):
    add_heading(doc, "Sign-Off & Distribution", 1)

    # Audit team table
    p = doc.add_paragraph()
    _run(p.add_run("Audit Team"), 11, bold=True, color=DARK_NAVY)
    _para_spacing(p, 0, 3)
    header = ["Role", "Name / Title", "Signature", "Date"]
    rows = [[r.get("role",""), r.get("name",""), "_________________", report_date]
            for r in audit_team]
    add_markdown_table(doc, header, rows)

    # Distribution list
    p2 = doc.add_paragraph()
    _run(p2.add_run("Distribution List"), 11, bold=True, color=DARK_NAVY)
    _para_spacing(p2, 4, 3)
    for d in distribution:
        add_bullet(doc, d)
    add_spacer(doc, 6)

    # Disclaimer
    disc = doc.add_paragraph()
    _para_spacing(disc, 4, 4)
    _run(disc.add_run(
        "This report is intended solely for the use of Wynn Resorts management and the Board Audit Committee. "
        "It is based on information available at the time of the audit and should not be relied upon for purposes "
        "other than those for which it was prepared. Distribution outside of the named recipients requires written "
        "approval from the Chief Audit Executive."
    ), 8, italic=True, color=MID_GREY)
    add_horizontal_rule(doc, GOLD, 1)
